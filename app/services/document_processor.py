"""
==============================================================
app/services/document_processor.py — Document Processing Service
==============================================================

WHY THIS FILE EXISTS:
    When a user uploads a PDF, Word doc, or text file, we need to:
    1. Read the raw bytes of the file
    2. Extract the text content (PDFs are not plain text!)
    3. Clean the text
    4. Split it into chunks (pieces of ~500 characters)

    This file handles ALL document reading and text extraction.

WHAT IS DOCUMENT PROCESSING IN RAG?
    The RAG (Retrieval-Augmented Generation) pipeline starts here.

    Raw File → Extract Text → Clean Text → Split into Chunks → ...
                                                              (continues in vector_store.py)

    Think of it like processing a book:
    - "Read the book" = Extract text from PDF
    - "Summarize each chapter" = Create chunks
    - "File each summary" = Store in vector DB

BEGINNER TIP - WHY NOT JUST SEND THE WHOLE DOCUMENT TO THE LLM?
    1. Token limits: GPT-4/Gemini can only process ~8,000-32,000 tokens
       at a time. A 100-page PDF might have 100,000 tokens!
    2. Cost: More tokens = more money. Relevant chunks = cheaper.
    3. Quality: LLMs perform BETTER with focused, relevant context
       than with huge documents full of irrelevant content.

HOW IT CONNECTS:
    document_processor.py ← called by documents_router.py
    document_processor.py → returns chunks to documents_router.py
    documents_router.py → passes chunks to vector_store.py
==============================================================
"""

import os
import io
from pathlib import Path
from typing import List, Tuple, Optional

import PyPDF2
import docx
import aiofiles

from app.utils.logger import logger
from app.utils.helpers import clean_text, chunk_text_by_sentences, format_file_size
from app.config.config import settings


class DocumentProcessor:
    """
    Service for processing uploaded documents into text chunks.

    Supports: PDF (.pdf), Plain Text (.txt), Word Documents (.docx)

    DESIGN PATTERN: Service Layer
        This class has no state (no instance variables) — all methods
        just take input and return output. This makes it easy to test.

    FLOW:
        1. Receive file bytes + filename
        2. Detect file type from extension
        3. Extract text using the appropriate extractor
        4. Clean and chunk the text
        5. Return list of (chunk_text, chunk_metadata) tuples
    """

    def __init__(self):
        """Initialize the document processor."""
        logger.info("DocumentProcessor initialized")

    async def process_document(
        self,
        file_content: bytes,
        filename: str,
        document_id: str
    ) -> Tuple[List[str], dict]:
        """
        Main entry point: process a document into text chunks.

        This is an ASYNC function because reading large files can
        be slow. Making it async means other requests aren't blocked
        while we wait for file processing to complete.

        WHAT DOES ASYNC MEAN?
            In Python, most code runs sequentially (line by line).
            Async code allows the server to handle OTHER requests
            while waiting for slow operations (file I/O, API calls).

            Imagine a restaurant:
            - Sync: Waiter takes order, stands at kitchen, waits, delivers
            - Async: Waiter takes order, gives to kitchen, serves other tables,
                     comes back when food is ready

        Args:
            file_content: Raw bytes of the uploaded file
            filename:     Original filename (to detect file type)
            document_id:  Unique ID for this document

        Returns:
            Tuple of:
            - List of text chunks (strings)
            - Metadata dict (filename, type, size, chunk count)

        Raises:
            ValueError: If file type is not supported
            Exception:  If text extraction fails
        """
        logger.info(f"Processing document: {filename} (ID: {document_id})")

        # Detect file type from extension
        file_extension = Path(filename).suffix.lower()
        file_size = len(file_content)

        logger.info(f"File type: {file_extension} | Size: {format_file_size(file_size)}")

        # Extract text based on file type
        if file_extension == ".pdf":
            raw_text = await self._extract_from_pdf(file_content, filename)
        elif file_extension == ".txt":
            raw_text = await self._extract_from_txt(file_content, filename)
        elif file_extension == ".docx":
            raw_text = await self._extract_from_docx(file_content, filename)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}. Supported: .pdf, .txt, .docx")

        # Validate we got some text
        if not raw_text or not raw_text.strip():
            raise ValueError(f"No text could be extracted from {filename}. The file may be empty or image-only.")

        logger.info(f"Extracted {len(raw_text)} characters from {filename}")

        # Split into chunks
        # chunk_size and chunk_overlap are from settings (config.py)
        chunks = chunk_text_by_sentences(
            raw_text,
            chunk_size=settings.chunk_size,
            overlap=settings.chunk_overlap
        )

        if not chunks:
            raise ValueError(f"No chunks created from {filename}. The document may be too short.")

        logger.info(f"Created {len(chunks)} chunks from {filename}")

        # Build metadata for this document
        # This metadata is stored alongside embeddings in ChromaDB
        metadata = {
            "document_id": document_id,
            "filename": filename,
            "file_type": file_extension,
            "file_size": file_size,
            "file_size_human": format_file_size(file_size),
            "total_chunks": len(chunks),
            "total_characters": len(raw_text),
        }

        return chunks, metadata

    async def _extract_from_pdf(self, file_content: bytes, filename: str) -> str:
        """
        Extract text from a PDF file.

        HOW PDF TEXT EXTRACTION WORKS:
            PDFs store text in a special binary format. PyPDF2 reads
            the binary format and extracts the raw text from each page.

            LIMITATION: Image-only PDFs (scanned documents) have NO
            extractable text — they're just pictures! For those, you'd
            need OCR (Optical Character Recognition) like Tesseract.

        Args:
            file_content: Raw PDF bytes
            filename:     Filename for logging

        Returns:
            Extracted text string from all pages
        """
        logger.debug(f"Extracting text from PDF: {filename}")

        try:
            # Create a file-like object from bytes
            # io.BytesIO wraps bytes in a file-like interface
            # PyPDF2.PdfReader expects a file object, not raw bytes
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            # Check if PDF is encrypted (password-protected)
            if pdf_reader.is_encrypted:
                raise ValueError(f"PDF '{filename}' is password-protected. Please provide an unprotected PDF.")

            total_pages = len(pdf_reader.pages)
            logger.info(f"PDF has {total_pages} pages")

            # Extract text from each page
            all_text_parts = []
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        all_text_parts.append(f"[Page {page_num}]\n{page_text}")
                except Exception as e:
                    # Skip problematic pages but continue processing
                    logger.warning(f"Could not extract text from page {page_num}: {e}")
                    continue

            if not all_text_parts:
                raise ValueError(
                    f"No text found in PDF '{filename}'. "
                    "This may be a scanned/image-only PDF. "
                    "Please use a text-based PDF."
                )

            full_text = "\n\n".join(all_text_parts)
            logger.debug(f"Extracted {len(full_text)} chars from {total_pages} pages")
            return full_text

        except Exception as e:
            logger.error(f"PDF extraction failed for {filename}: {e}")
            raise

    async def _extract_from_txt(self, file_content: bytes, filename: str) -> str:
        """
        Extract text from a plain text file.

        Text files are the simplest case — we just decode the bytes
        to a string. But we need to handle different encodings:
        - UTF-8: Most common, supports Unicode (emojis, accents, etc.)
        - Latin-1: Older format, common in legacy systems
        - ASCII: Very basic, only English characters

        Args:
            file_content: Raw text file bytes
            filename:     Filename for logging

        Returns:
            Decoded text string
        """
        logger.debug(f"Extracting text from TXT: {filename}")

        # Try different encodings in order of preference
        encodings_to_try = ["utf-8", "utf-8-sig", "latin-1", "ascii"]

        for encoding in encodings_to_try:
            try:
                text = file_content.decode(encoding)
                logger.debug(f"Successfully decoded with {encoding} encoding")
                return text
            except UnicodeDecodeError:
                continue

        # If all encodings fail, use utf-8 with error replacement
        logger.warning(f"Could not cleanly decode {filename}, using lossy UTF-8 decode")
        return file_content.decode("utf-8", errors="replace")

    async def _extract_from_docx(self, file_content: bytes, filename: str) -> str:
        """
        Extract text from a Microsoft Word (.docx) file.

        HOW DOCX EXTRACTION WORKS:
            .docx files are actually ZIP archives containing XML files.
            python-docx opens this ZIP and parses the XML to extract text.

            A Word document has:
            - Paragraphs: Main text blocks
            - Tables: Grid data
            - Headers/Footers: Page headers and footers

            We extract all paragraphs and tables.

        Args:
            file_content: Raw .docx bytes
            filename:     Filename for logging

        Returns:
            Extracted text string
        """
        logger.debug(f"Extracting text from DOCX: {filename}")

        try:
            # Create file-like object from bytes
            docx_file = io.BytesIO(file_content)
            doc = docx.Document(docx_file)

            text_parts = []

            # Extract text from all paragraphs
            # Each paragraph is a separate text block in Word
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            # Tables have rows and cells — we extract each cell's text
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_texts.append(cell.text.strip())
                    if row_texts:
                        # Join cells with " | " to represent table structure
                        text_parts.append(" | ".join(row_texts))

            if not text_parts:
                raise ValueError(f"No text found in Word document '{filename}'")

            full_text = "\n".join(text_parts)
            logger.debug(f"Extracted {len(full_text)} chars from DOCX")
            return full_text

        except Exception as e:
            logger.error(f"DOCX extraction failed for {filename}: {e}")
            raise

    def validate_file_size(self, file_size: int) -> bool:
        """
        Check if a file is within the allowed size limit.

        Args:
            file_size: File size in bytes

        Returns:
            True if within limit, False if too large

        Example:
            # settings.max_file_size = 52428800 (50MB)
            validate_file_size(1024 * 1024)       → True  (1MB is fine)
            validate_file_size(100 * 1024 * 1024) → False (100MB is too large)
        """
        return file_size <= settings.max_file_size


# ==============================================================
# SINGLETON INSTANCE
# ==============================================================

document_processor = DocumentProcessor()
